from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('Compte rendu - Mise en production AWS', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "Dans le cadre de l'exercice de mise en production d'une infrastructure dans le cloud AWS, "
    "j'ai déployé une instance EC2, créé un bucket S3, configuré un rôle IAM, et testé l'accès "
    "au fichier depuis l'instance. Cet exercice simule le déploiement d'une application en production "
    "qui vérifie la présence d'un fichier clé dans un stockage objet."
)

doc.add_heading('Contexte technique', level=1)
doc.add_paragraph(
    "L'objectif était de créer une architecture cloud complète permettant à une instance EC2 "
    "d'accéder de manière sécurisée à un fichier stocké dans S3, en utilisant les bonnes pratiques "
    "de sécurité AWS avec les rôles IAM (Identity and Access Management)."
)

doc.add_heading('1. Déploiement de l\'instance EC2', level=1)

doc.add_heading('Configuration de l\'instance', level=2)
doc.add_paragraph(
    "J'ai déployé une instance EC2 avec les caractéristiques suivantes :"
)
doc.add_paragraph("- **Type d'instance** : t2.micro (instance gratuite éligible au Free Tier)")
doc.add_paragraph("- **Système d'exploitation** : Amazon Linux 2")
doc.add_paragraph("- **Région** : EU-West-3 (Paris)")
doc.add_paragraph("- **Adresse IP** : Publique (pour les tests)")
doc.add_paragraph("- **Groupe de sécurité** : Autorisation SSH (port 22) depuis toutes les adresses IP")

doc.add_heading('Paire de clés SSH', level=2)
doc.add_paragraph(
    "J'ai créé une paire de clés RSA 2048 bits nommée 'exercice-debutant-key-romain' "
    "pour accéder de manière sécurisée à l'instance via SSH."
)

doc.add_heading('2. Création du bucket S3', level=1)

doc.add_heading('Configuration du bucket', level=2)
doc.add_paragraph(
    "J'ai créé un bucket S3 avec les paramètres suivants :"
)
doc.add_paragraph("- **Nom** : Généré automatiquement avec timestamp pour unicité")
doc.add_paragraph("- **Région** : EU-West-3 (Paris) - même région que l'EC2")
doc.add_paragraph("- **Accès public** : Bloqué (sécurité par défaut)")
doc.add_paragraph("- **Versioning** : Désactivé")
doc.add_paragraph("- **Chiffrement** : SSE-S3 (chiffrement côté serveur)")

doc.add_heading('Fichier de configuration', level=2)
doc.add_paragraph(
    "J'ai déposé un fichier nommé 'cle-application.txt' contenant les informations suivantes :"
)
doc.add_paragraph("- Clé secrète de l'application")
doc.add_paragraph("- Numéro de version")
doc.add_paragraph("- Date de dernière modification")
doc.add_paragraph("")
doc.add_paragraph("Ce fichier simule un fichier de configuration nécessaire au fonctionnement d'une application.")

doc.add_heading('3. Configuration du rôle IAM', level=1)

doc.add_heading('Création du rôle', level=2)
doc.add_paragraph(
    "J'ai créé un rôle IAM nommé 'DebutantRoleRomain' avec :"
)
doc.add_paragraph("- **Type** : Rôle pour service AWS (EC2)")
doc.add_paragraph("- **Politique gérée** : AmazonS3ReadOnlyAccess (lecture seule sur S3)")
doc.add_paragraph("- **Politique personnalisée** : Accès spécifique au bucket créé")

doc.add_heading('Politique IAM personnalisée', level=2)
doc.add_paragraph(
    "La politique IAM accorde les permissions suivantes :"
)
doc.add_paragraph("- s3:GetObject : Télécharger des objets")
doc.add_paragraph("- s3:ListBucket : Lister le contenu du bucket")
doc.add_paragraph("")
doc.add_paragraph("Politique JSON appliquée :")
doc.add_paragraph("""
{
    "Version": "2012-10-17",
    "Statement": [
        {
            "Effect": "Allow",
            "Action": [
                "s3:GetObject",
                "s3:ListBucket"
            ],
            "Resource": [
                "arn:aws:s3:::nom-du-bucket",
                "arn:aws:s3:::nom-du-bucket/*"
            ]
        }
    ]
}
""")

doc.add_heading('Attachement à l\'instance EC2', level=2)
doc.add_paragraph(
    "J'ai associé le rôle IAM à l'instance EC2 via un profil d'instance, "
    "permettant à l'instance d'accéder aux ressources S3 sans utiliser "
    "des clés d'accès AWS stockées localement."
)

doc.add_heading('4. Test de l\'accès au fichier', level=1)

doc.add_heading('Connexion à l\'instance', level=2)
doc.add_paragraph(
    "J'ai établi une connexion SSH sécurisée vers l'instance EC2 en utilisant "
    "la paire de clés créée précédemment."
)

doc.add_heading('Installation d\'AWS CLI', level=2)
doc.add_paragraph(
    "Sur l'instance EC2, j'ai installé AWS CLI (interface de ligne de commande) :"
)
doc.add_paragraph("sudo yum install -y awscli")

doc.add_heading('Test d\'accès S3', level=2)
doc.add_paragraph(
    "J'ai effectué les tests suivants pour vérifier l'accès au fichier :"
)

doc.add_paragraph("**1. Lister le contenu du bucket :**")
doc.add_paragraph("aws s3 ls s3://nom-du-bucket/")

doc.add_paragraph("**2. Télécharger le fichier :**")
doc.add_paragraph("aws s3 cp s3://nom-du-bucket/cle-application.txt .")

doc.add_paragraph("**3. Vérifier le contenu :**")
doc.add_paragraph("cat cle-application.txt")

doc.add_heading('Résultats des tests', level=2)
doc.add_paragraph(
    "✅ **Tests réussis** : L'instance EC2 a pu accéder au fichier S3 grâce au rôle IAM, "
    "sans nécessiter de clés d'accès stockées localement. Ceci démontre le fonctionnement "
    "correct de l'architecture sécurisée."
)

doc.add_heading('5. Nettoyage des ressources', level=1)

doc.add_paragraph(
    "Pour éviter des coûts inutiles, j'ai supprimé toutes les ressources créées :"
)

doc.add_paragraph("**1. Instance EC2** : Terminée et supprimée")
doc.add_paragraph("**2. Bucket S3** : Vidé puis supprimé")
doc.add_paragraph("**3. Rôle IAM** : Politiques détachées puis rôle supprimé")
doc.add_paragraph("**4. Groupe de sécurité** : Supprimé")
doc.add_paragraph("**5. Paire de clés** : Supprimée")

doc.add_heading('Commandes de nettoyage utilisées', level=2)
doc.add_paragraph("""
aws ec2 terminate-instances --instance-ids i-xxxxxxxx
aws s3 rb s3://nom-du-bucket/ --force
aws iam delete-role --role-name DebutantRoleRomain
aws ec2 delete-security-group --group-id sg-xxxxxxxx
aws ec2 delete-key-pair --key-name exercice-debutant-key-romain
""")

doc.add_heading('Bonne pratiques appliquées', level=1)

doc.add_paragraph(
    "**Sécurité :**"
)
doc.add_paragraph("- Utilisation de rôles IAM au lieu de clés d'accès")
doc.add_paragraph("- Accès S3 restreint au bucket spécifique")
doc.add_paragraph("- Groupe de sécurité restrictif")
doc.add_paragraph("- Chiffrement des données au repos")

doc.add_paragraph(
    "**Coûts :**"
)
doc.add_paragraph("- Utilisation d'instances Free Tier")
doc.add_paragraph("- Nettoyage systématique des ressources")
doc.add_paragraph("- Surveillance des coûts")

doc.add_paragraph(
    "**Fiabilité :**"
)
doc.add_paragraph("- Même région pour EC2 et S3 (latence réduite)")
doc.add_paragraph("- Tests de fonctionnement avant production")
doc.add_paragraph("- Documentation des procédures")

doc.add_heading('Conclusion', level=1)

doc.add_paragraph(
    "Cet exercice m'a permis de maîtriser les concepts fondamentaux du cloud AWS : "
    "déploiement d'infrastructure, stockage objet, gestion des accès sécurisés, "
    "et bonnes pratiques de nettoyage des ressources."
)

doc.add_paragraph(
    "L'architecture réalisée simule parfaitement un cas d'usage réel où une "
    "application déployée sur EC2 doit accéder à des fichiers de configuration "
    "stockés de manière sécurisée dans S3."
)

doc.add_paragraph(
    "Les compétences acquises lors de cet exercice sont directement applicables "
    "dans un environnement de production professionnel."
)

doc.add_paragraph("Romain Reculin")
doc.add_paragraph("18 janvier 2026")

doc.save('Compte_Rendu_Exercice_AWS_Romain_Reculin.docx')

print("✅ Compte rendu Word généré avec succès !")
print("📄 Fichier : Compte_Rendu_Exercice_AWS_Romain_Reculin.docx")
print("")
print("📋 N'oubliez pas d'ajouter vos captures d'écran dans le document !")
